#!/usr/bin/env python3
"""Extracts Resume_Rahman.docx and writes README.md"""

from pathlib import Path
from docx import Document
from docx.table import Table

STYLE_MAP = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
}

def extract_table(table):
    lines = []
    rows = table.rows
    if not rows:
        return ""
    # header row
    headers = [cell.text.strip() for cell in rows[0].cells]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows[1:]:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def docx_to_markdown(path: Path) -> str:
    doc = Document(str(path))
    md_lines = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            # find matching paragraph
            for para in doc.paragraphs:
                if para._element is block:
                    style = para.style.name
                    text = para.text.strip()
                    if not text:
                        md_lines.append("")
                    elif style in STYLE_MAP:
                        md_lines.append(f"{STYLE_MAP[style]} {text}")
                    elif style == "List Bullet" or style.startswith("List Bullet"):
                        md_lines.append(f"- {text}")
                    elif style == "List Number" or style.startswith("List Number"):
                        md_lines.append(f"1. {text}")
                    else:
                        md_lines.append(text)
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is block:
                    md_lines.append(extract_table(table))
                    md_lines.append("")
                    break

    return "\n".join(md_lines)


if __name__ == "__main__":
    src = Path(__file__).parent / "Resume_Rahman.docx"
    out = Path(__file__).parent / "Resume.md"
    content = docx_to_markdown(src)
    out.write_text(content, encoding="utf-8")
    print(f"Written to {out}")
