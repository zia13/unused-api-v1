#!/usr/bin/env python3
"""
build_resume_docx.py
────────────────────
Converts Resume.md → Resume.docx with professional resume formatting.

Install deps (once):
    pip install python-docx

Run:
    python build_resume_docx.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("ERROR: python-docx is not installed.\nRun: pip install python-docx")

MD_FILE  = Path(__file__).parent / "Resume.md"
OUT_FILE = Path(__file__).parent / "Resume_Rahman.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x6D, 0xB4)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)
MID_GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# Section headers that mark top-level resume sections
SECTION_HEADERS = {
    "SUMMARY", "CORE COMPETENCIES", "RECENT TOOLS",
    "RECENT EXPERIENCE", "EDUCATION", "CERTIFICATIONS",
}

# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_para_border_bottom(para, colour="CCCCCC", sz="6"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), colour)
    pBdr.append(bot)
    pPr.append(pBdr)


def remove_para_spacing(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def set_para_shading(para, hex_colour: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    pPr.append(shd)


def set_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BBBBBB")
        tblBdr.append(el)
    tblPr.append(tblBdr)


def set_col_width(table, col_idx, width_cm):
    """Set preferred width on every cell in a column."""
    for row in table.rows:
        tc   = row.cells[col_idx]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(int(width_cm * 567)))   # 1 cm ≈ 567 twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


# ── Paragraph helpers ──────────────────────────────────────────────────────────

def add_section_header(doc, text: str):
    """Shaded banner with white bold text — mirrors typical resume section bars."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, "1F497D")
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = WHITE
    return p


def add_name_block(doc, name: str, title: str):
    """Centered name + title at the top of the document."""
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_para_spacing(p_name)
    r = p_name.add_run(name)
    r.bold           = True
    r.font.size      = Pt(22)
    r.font.color.rgb = DARK_BLUE

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after  = Pt(8)
    r2 = p_title.add_run(title)
    r2.font.size      = Pt(12)
    r2.font.color.rgb = MID_BLUE
    r2.bold           = True

    # Thin rule under the name block
    set_para_border_bottom(p_title, colour="1F497D", sz="12")


def add_bullet(doc, text: str, indent_level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 + 0.2 * indent_level)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(1)
    inline_styled(p, text)
    return p


def inline_styled(para, text: str):
    """Add runs with **bold**, *italic*, `code` inline support."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("**"):
            r = para.add_run(m.group(2)); r.bold = True
        elif raw.startswith("*"):
            r = para.add_run(m.group(3)); r.italic = True
        else:
            r = para.add_run(m.group(4))
            r.font.name = "Courier New"; r.font.size = Pt(9.5)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def add_experience_header(doc, company_dates: str, location: str, job_title: str):
    """Two-column table row: company+dates | nothing; then location; then title."""
    # Company line with dates on right
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    # Remove all borders so it looks invisible
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")
        tblBdr.append(el)
    tblPr.append(tblBdr)

    left_cell, right_cell = table.rows[0].cells

    # Parse "Company Name   DATES" — split on 2+ spaces or a tab
    parts = re.split(r"\t|  {2,}", company_dates, maxsplit=1)
    company = parts[0].strip()
    dates   = parts[1].strip() if len(parts) > 1 else ""

    lp = left_cell.paragraphs[0]
    lr = lp.add_run(company)
    lr.bold = True; lr.font.size = Pt(10.5); lr.font.color.rgb = DARK_BLUE

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(dates)
    rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = DARK_GREY

    for cell in (left_cell, right_cell):
        cell.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(0)

    # Location
    if location:
        pl = doc.add_paragraph()
        pl.paragraph_format.space_before = Pt(0)
        pl.paragraph_format.space_after  = Pt(0)
        rl = pl.add_run(location)
        rl.font.size = Pt(9.5); rl.font.color.rgb = MID_GREY; rl.italic = True

    # Job title
    pt = doc.add_paragraph()
    pt.paragraph_format.space_before = Pt(1)
    pt.paragraph_format.space_after  = Pt(2)
    rt = pt.add_run(job_title)
    rt.bold = True; rt.font.size = Pt(10.5); rt.font.color.rgb = DARK_GREY


def add_tools_table(doc, rows: list[tuple[str, str]]):
    """Two-column shaded table for the RECENT TOOLS section."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_col_width(table, 0, 4.2)
    set_col_width(table, 1, 12.1)

    for i, (key, value) in enumerate(rows):
        key_cell, val_cell = table.rows[i].cells
        fill = "EAF0F8" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(key_cell, fill)
        set_cell_bg(val_cell, fill)

        kp = key_cell.paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True; kr.font.size = Pt(9.5); kr.font.color.rgb = DARK_BLUE

        vp = val_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(9.5); vr.font.color.rgb = DARK_GREY

        for cell in (key_cell, val_cell):
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Main builder ───────────────────────────────────────────────────────────────

def build_resume(md_path: Path, out_path: Path):
    doc  = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    lines   = md_path.read_text(encoding="utf-8").splitlines()
    n       = len(lines)
    i       = 0

    # ── State machine ──────────────────────────────────────────────────────────
    current_section = None

    # Detect name / title at very top
    name_line  = lines[0].strip() if n > 0 else ""
    title_line = lines[1].strip() if n > 1 else ""
    if name_line and not name_line.startswith("#"):
        add_name_block(doc, name_line, title_line)
        i = 2

    # Track experience-block state
    in_experience    = False
    exp_company_line = ""
    exp_location     = ""
    exp_title        = ""
    exp_bullets      = []   # list of (text, is_environment)

    def flush_experience():
        nonlocal in_experience, exp_company_line, exp_location, exp_title, exp_bullets
        if not in_experience:
            return
        add_experience_header(doc, exp_company_line, exp_location, exp_title)
        for (bullet_text, is_env) in exp_bullets:
            if is_env:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(4)
                p.paragraph_format.left_indent  = Inches(0.1)
                r = p.add_run(bullet_text)
                r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MID_GREY
            else:
                add_bullet(doc, bullet_text)
        in_experience    = False
        exp_company_line = ""
        exp_location     = ""
        exp_title        = ""
        exp_bullets      = []

    # Tools section accumulator
    tools_rows: list[tuple[str, str]] = []

    def flush_tools():
        nonlocal tools_rows
        if tools_rows:
            add_tools_table(doc, tools_rows)
            tools_rows = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── Section header detection ───────────────────────────────────────────
        if stripped in SECTION_HEADERS:
            flush_experience()
            flush_tools()
            current_section = stripped
            add_section_header(doc, stripped)
            i += 1
            continue

        # Skip blank lines (but flush tools if we hit blank while in tools section)
        if stripped == "":
            if current_section == "RECENT TOOLS" and tools_rows:
                flush_tools()
            i += 1
            continue

        # ── RECENT TOOLS ──────────────────────────────────────────────────────
        if current_section == "RECENT TOOLS":
            # Lines like "Back End\tJava, ..."  or  "Back End   Java, ..."
            tab_match = re.match(r"^(.+?)\t+(.+)$", stripped)
            sp_match  = re.match(r"^([A-Za-z &/\-]+?)\s{3,}(.+)$", stripped) if not tab_match else None
            if tab_match:
                tools_rows.append((tab_match.group(1).strip(), tab_match.group(2).strip()))
                i += 1
                continue
            elif sp_match:
                tools_rows.append((sp_match.group(1).strip(), sp_match.group(2).strip()))
                i += 1
                continue
            else:
                # Continuation or standalone: append to last row value
                if tools_rows:
                    key, val = tools_rows[-1]
                    tools_rows[-1] = (key, val + " " + stripped)
                i += 1
                continue

        # ── RECENT EXPERIENCE ─────────────────────────────────────────────────
        if current_section == "RECENT EXPERIENCE":
            # Detect company line: has a date-like pattern  e.g.  "Company  JAN 2020 – Present"
            company_date_match = re.match(
                r"^(.+?(?:Inc|Ltd|Lab|Technologies|Canada|IFIL|Vivasoft|Kona|Apurba|Object|Ipvision|Mphasis|Data).+?)\s{2,}(.+(?:\d{4}|Present).*)$",
                stripped, re.IGNORECASE
            )
            # Also try tab-separated
            if not company_date_match:
                company_date_match = re.match(
                    r"^(.+?(?:Inc|Ltd|Lab|Technologies|Canada|IFIL|Vivasoft|Kona|Apurba|Object|Ipvision|Mphasis|Data).+?)\t(.+)$",
                    stripped, re.IGNORECASE
                )
            # Fallback: any line with  2+ spaces followed by a date range
            if not company_date_match:
                company_date_match = re.match(
                    r"^(.+?)\s{2,}(\w{3}\s+\d{4}.+(?:\d{4}|Present).*)$",
                    stripped, re.IGNORECASE
                )

            # Location line: "City, Country" pattern (short, no verb words)
            location_match = re.match(r"^([A-Za-z ]+,\s*[A-Za-z ]+)$", stripped)

            # Job title line: short line after location that isn't a bullet and has no colon
            is_job_title = (
                in_experience and exp_title == ""
                and not stripped.startswith("-")
                and not stripped.lower().startswith("environment")
                and len(stripped) < 80
                and ":" not in stripped
                and not company_date_match
            )

            if company_date_match:
                flush_experience()
                in_experience    = True
                exp_company_line = stripped
                exp_location     = ""
                exp_title        = ""
                exp_bullets      = []
                i += 1
                continue

            if in_experience and location_match and exp_title == "":
                exp_location = stripped
                i += 1
                continue

            if in_experience and is_job_title and exp_location != "":
                exp_title = stripped
                i += 1
                continue

            # Bullet points
            bullet_match = re.match(r"^[-•]\s+(.*)", stripped)
            if bullet_match and in_experience:
                exp_bullets.append((bullet_match.group(1), False))
                i += 1
                continue

            # Plain sentences inside experience (non-bullet responsibilities)
            if in_experience and exp_title and not bullet_match:
                if stripped.lower().startswith("environment:") or stripped.lower().startswith("environment "):
                    exp_bullets.append((stripped, True))
                else:
                    exp_bullets.append((stripped, False))
                i += 1
                continue

            # Fallback: plain paragraph
            flush_experience()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            inline_styled(p, stripped)
            i += 1
            continue

        # ── SUMMARY ───────────────────────────────────────────────────────────
        if current_section == "SUMMARY":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(4)
            inline_styled(p, stripped)
            i += 1
            continue

        # ── CORE COMPETENCIES ─────────────────────────────────────────────────
        if current_section == "CORE COMPETENCIES":
            # Lines like "Architecture & Development: Expert in ..."
            colon_match = re.match(r"^([^:]+):\s*(.+)$", stripped)
            if colon_match:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(2)
                r_key = p.add_run(colon_match.group(1) + ": ")
                r_key.bold = True; r_key.font.color.rgb = DARK_BLUE
                p.add_run(colon_match.group(2))
            else:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                inline_styled(p, stripped)
            i += 1
            continue

        # ── EDUCATION ─────────────────────────────────────────────────────────
        if current_section == "EDUCATION":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            inline_styled(p, stripped)
            i += 1
            continue

        # ── CERTIFICATIONS ────────────────────────────────────────────────────
        if current_section == "CERTIFICATIONS":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            inline_styled(p, stripped)
            i += 1
            continue

        # ── Default ───────────────────────────────────────────────────────────
        p = doc.add_paragraph()
        inline_styled(p, stripped)
        i += 1

    # Flush any remaining state
    flush_experience()
    flush_tools()

    doc.save(str(out_path))
    print(f"✅  Saved → {out_path}")


if __name__ == "__main__":
    build_resume(MD_FILE, OUT_FILE)
