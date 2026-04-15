#!/usr/bin/env python3
"""
build_docx.py
─────────────
Converts aws-api-gateway-unused-api-cleanup.md → aws-api-gateway-unused-api-cleanup.docx
using python-docx.

Install deps (once):
    pip install python-docx

Run:
    python build_docx.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    sys.exit(
        "ERROR: python-docx is not installed.\n"
        "Run:  pip install python-docx\n"
        "Then re-run this script."
    )

MD_FILE  = Path(__file__).parent / "aws-api-gateway-unused-api-cleanup.md"
OUT_FILE = Path(__file__).parent / "aws-api-gateway-unused-api-cleanup.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
BLUE        = RGBColor(0x1F, 0x49, 0x7D)   # heading blue
LIGHT_BLUE  = RGBColor(0xD6, 0xE4, 0xF0)   # table header fill
CODE_BG     = RGBColor(0xF4, 0xF4, 0xF4)   # code block background
CODE_FG     = RGBColor(0x24, 0x29, 0x2E)   # code text colour
GREY_BORDER = RGBColor(0xCC, 0xCC, 0xCC)


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_table_borders(table):
    """Add thin borders around every cell in a table."""
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_horizontal_rule(doc):
    """Insert a light grey horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


# ── Style helpers ──────────────────────────────────────────────────────────────

def style_heading(para, level: int):
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.font.color.rgb = BLUE
    run.font.bold      = True
    run.font.size      = Pt([0, 20, 16, 14, 13, 12][min(level, 5)])


def add_heading(doc, text: str, level: int):
    h = doc.add_heading(text, level=level)
    style_heading(h, level)
    return h


def add_code_block(doc, code: str):
    """Add a shaded code paragraph with monospace font."""
    # Split into lines; each line is its own paragraph so shading tiles correctly
    lines = code.split("\n")
    for i, line in enumerate(lines):
        p   = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = CODE_FG
        # Shade the paragraph
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F4F4F4")
        pPr.append(shd)
        # No space between lines inside a block
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
    # Small gap after block
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table_from_md(doc, header_row: list[str], data_rows: list[list[str]]):
    cols  = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"
    set_table_borders(table)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header_row):
        hdr_cells[i].text = text.strip().strip("*")
        set_cell_bg(hdr_cells[i], "D6E4F0")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold       = True
            run.font.color.rgb  = BLUE
            run.font.size       = Pt(10)

    # Data rows
    for r, row in enumerate(data_rows):
        cells = table.rows[r + 1].cells
        for c, text in enumerate(row):
            cells[c].text = text.strip()
            for run in cells[c].paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


def inline_bold_italic(para, text: str):
    """
    Parse inline **bold**, *italic*, and `code` spans and add styled runs.
    """
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > last:
            para.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("**"):
            r = para.add_run(m.group(2))
            r.bold = True
        elif raw.startswith("*"):
            r = para.add_run(m.group(3))
            r.italic = True
        else:  # backtick code
            r = para.add_run(m.group(4))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
            r.font.color.rgb = CODE_FG
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


# ── Main parser / builder ──────────────────────────────────────────────────────

def parse_md_table(lines: list[str], start: int):
    """
    Parse a GFM table starting at `start`.
    Returns (header_cells, data_rows, next_line_index).
    """
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    # skip separator row
    i       = start + 2
    rows    = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    return header, rows, i


def build_docx(md_path: Path, out_path: Path):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    lines      = md_path.read_text(encoding="utf-8").splitlines()
    i          = 0
    in_code    = False
    code_lines = []
    code_lang  = ""

    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ─────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code   = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── ATX headings ──────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,5})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text  = heading_match.group(2).strip()
            # Strip markdown link anchors like [text](#anchor)
            text  = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
            add_heading(doc, text, level)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()) or re.match(r"^\*{3,}$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── GFM table ─────────────────────────────────────────────────────────
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-\|:]+\|\s*$", lines[i + 1]):
            header, rows, next_i = parse_md_table(lines, i)
            add_table_from_md(doc, header, rows)
            i = next_i
            continue

        # ── Bullet / unordered list ───────────────────────────────────────────
        ul_match = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if ul_match:
            indent = len(ul_match.group(1)) // 2
            text   = ul_match.group(2)
            p      = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            inline_bold_italic(p, text)
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────────────────────────
        ol_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if ol_match:
            indent = len(ol_match.group(1)) // 2
            text   = ol_match.group(2)
            p      = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            inline_bold_italic(p, text)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        bq_match = re.match(r"^>\s*(.*)", line)
        if bq_match:
            text = bq_match.group(1)
            p    = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.italic          = True
            run.font.color.rgb  = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        # Collect wrapped lines into one paragraph
        para_lines = [line]
        while i + 1 < len(lines) and lines[i + 1].strip() != "" \
                and not re.match(r"^#{1,5}\s", lines[i + 1]) \
                and not lines[i + 1].strip().startswith("|") \
                and not lines[i + 1].strip().startswith("```") \
                and not re.match(r"^[-*+]\s", lines[i + 1]) \
                and not re.match(r"^\d+\.\s", lines[i + 1]) \
                and not re.match(r"^>\s", lines[i + 1]):
            i += 1
            para_lines.append(lines[i])
        full_text = " ".join(para_lines)
        # Strip markdown image/link syntax
        full_text = re.sub(r"!\[.*?\]\(.*?\)", "", full_text)
        full_text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", full_text)
        if full_text.strip():
            p = doc.add_paragraph()
            inline_bold_italic(p, full_text)

        i += 1

    doc.save(str(out_path))
    print(f"✅  Saved: {out_path}")


if __name__ == "__main__":
    build_docx(MD_FILE, OUT_FILE)
